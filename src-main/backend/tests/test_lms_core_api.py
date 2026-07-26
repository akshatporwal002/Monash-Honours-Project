from __future__ import annotations

from collections.abc import Generator
from datetime import UTC, datetime, timedelta
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from app.api.routes.lms import get_lms_material_storage
from app.db.base import Base
from app.db.session import create_db_engine, create_session_factory, get_db
from app.main import create_app
from app.models import (
    AttemptStatus,
    Course,
    CourseState,
    LearningEvent,
    LearningEventType,
    LearningMaterial,
    LearningTask,
    MaterialChunk,
    MaterialIndexStatus,
    PlatformAuditEvent,
    Recommendation,
    Reminder,
    SubmissionAttempt,
    TaskPointAward,
)
from app.services.lms import DEMO_PASSWORD, bootstrap_demo
from app.services.rag.storage import LocalFileStorage


@pytest.fixture
def lms_context(tmp_path: Path) -> Generator[tuple[TestClient, Session], None, None]:
    database_path = tmp_path / "lms.db"
    engine = create_db_engine(f"sqlite:///{database_path.as_posix()}")
    Base.metadata.create_all(engine)
    factory = create_session_factory(engine)
    session = factory()
    bootstrap_demo(session)

    app = create_app()
    app.dependency_overrides[get_db] = lambda: session
    app.dependency_overrides[get_lms_material_storage] = lambda: LocalFileStorage(
        tmp_path / "uploads",
        20 * 1024 * 1024,
    )
    try:
        with TestClient(app) as client:
            yield client, session
    finally:
        app.dependency_overrides.clear()
        session.close()
        Base.metadata.drop_all(engine)
        engine.dispose()


def login(client: TestClient, role: str) -> None:
    response = client.post(
        "/api/v1/auth/login",
        json={
            "email": f"{role}@quantumlearn.demo",
            "password": DEMO_PASSWORD,
        },
    )
    assert response.status_code == 200


def test_role_scoping_and_explicit_bootstrap(
    lms_context: tuple[TestClient, Session],
) -> None:
    client, session = lms_context
    assert client.get("/api/v1/courses").status_code == 401

    login(client, "student")
    assert (
        client.post(
            "/api/v1/courses",
            json={"code": "NOPE-1", "title": "Forbidden"},
        ).status_code
        == 403
    )
    client.post("/api/v1/auth/logout")

    login(client, "educator")
    courses = client.get("/api/v1/courses")
    assert courses.status_code == 200
    assert [course["code"] for course in courses.json()] == ["QL-101"]
    client.post("/api/v1/auth/logout")

    login(client, "admin")
    assert client.get("/api/v1/admin/settings").json()["at_risk_threshold"] == 70
    assert client.get("/api/v1/educator/dashboard").status_code == 403

    # The helper is idempotent and no read endpoint invokes it.
    users, course = bootstrap_demo(session)
    assert len(users) == 3
    assert session.scalar(select(func.count()).select_from(Course)) == 1
    assert course.code == "QL-101"


def test_generation_requires_authorised_course_sources_but_authored_tasks_do_not(
    lms_context: tuple[TestClient, Session],
) -> None:
    client, session = lms_context
    login(client, "educator")
    course = client.post(
        "/api/v1/courses",
        json={"code": "SRC-101", "title": "Source validation"},
    ).json()
    module = client.post(
        f"/api/v1/courses/{course['id']}/modules",
        json={"title": "Grounding", "description": "", "position": 1},
    ).json()
    outcome = client.post(
        f"/api/v1/modules/{module['id']}/outcomes",
        json={
            "title": "Explain grounding",
            "statement": "Explain why generated tasks cite authorised course sources.",
            "kind": "topic",
            "position": 1,
        },
    ).json()

    generated = client.post(
        f"/api/v1/courses/{course['id']}/generate-tasks",
        json={"learning_outcome_id": outcome["id"], "task_count": 3},
    )

    assert generated.status_code == 422
    assert generated.json()["detail"] == "No relevant course content was found."
    assert (
        session.scalar(
            select(func.count())
            .select_from(LearningTask)
            .where(LearningTask.course_id == course["id"])
        )
        == 0
    )

    authored = client.post(
        f"/api/v1/courses/{course['id']}/tasks",
        json={
            "module_id": module["id"],
            "learning_outcome_id": outcome["id"],
            "title": "Educator-authored check",
            "prompt": "Explain source grounding.",
            "instructions": "Answer in one sentence.",
            "task_type": "short_answer",
            "difficulty": "beginner",
            "position": 1,
            "expected_answer": "grounding",
        },
    )
    assert authored.status_code == 201
    assert authored.json()["source_references"] == []


def test_course_configuration_scaffolding_and_educator_scope(
    lms_context: tuple[TestClient, Session],
) -> None:
    client, session = lms_context
    login(client, "educator")
    course = client.post(
        "/api/v1/courses",
        json={
            "code": "QC-201",
            "title": "Quantum Algorithms",
            "description": "An MVP algorithms course.",
        },
    ).json()
    module = client.post(
        f"/api/v1/courses/{course['id']}/modules",
        json={"title": "Deutsch-Jozsa", "description": "", "position": 1},
    ).json()
    outcome = client.post(
        f"/api/v1/modules/{module['id']}/outcomes",
        json={
            "title": "Explain oracle behaviour",
            "statement": "Distinguish constant and balanced quantum oracles.",
            "kind": "topic",
            "position": 1,
        },
    ).json()
    material_response = client.post(
        f"/api/v1/courses/{course['id']}/materials/link",
        json={
            "source_url": "https://example.edu/deutsch-josza.pdf",
            "module_id": module["id"],
        },
    )
    assert material_response.status_code == 201
    material = material_response.json()
    stored_material = session.get(LearningMaterial, material["id"])
    assert stored_material is not None
    stored_material.indexing_status = MaterialIndexStatus.INDEXED
    stored_material.chunks.append(
        MaterialChunk(
            id="chunk-deutsch-jozsa",
            chunk_index=0,
            chunk_text=(
                "The Deutsch-Jozsa quantum oracle is constant when every input has "
                "the same output and balanced when half of the outputs differ."
            ),
            token_count=22,
            chunk_hash="sha256:deutsch-jozsa",
        )
    )
    session.commit()
    generated = client.post(
        f"/api/v1/courses/{course['id']}/generate-tasks",
        json={"learning_outcome_id": outcome["id"], "task_count": 6},
    )
    assert generated.status_code == 201
    tasks = generated.json()
    assert len(tasks) == 6
    assert {task["task_type"] for task in tasks} == {
        "multiple_choice",
        "multiple_answer",
        "short_answer",
        "code_explanation",
        "code_completion",
        "quantum_circuit",
    }
    assert tasks[0]["prerequisite_task_ids"] == []
    assert tasks[1]["prerequisite_task_ids"] == [tasks[0]["id"]]
    assert tasks[2]["prerequisite_task_ids"] == [tasks[1]["id"]]
    assert all(task["source_references"] == ["chunk-deutsch-jozsa"] for task in tasks)
    assert all("constant when every input" in task["prompt"] for task in tasks)

    missing_sources = client.patch(
        f"/api/v1/tasks/{tasks[0]['id']}",
        json={"source_references": []},
    )
    assert missing_sources.status_code == 422
    assert "at least one authorised" in missing_sources.json()["detail"]
    invented_source = client.patch(
        f"/api/v1/tasks/{tasks[0]['id']}",
        json={"source_references": ["invented-source"]},
    )
    assert invented_source.status_code == 422
    assert "in this course" in invented_source.json()["detail"]

    published = client.post(f"/api/v1/courses/{course['id']}/publish")
    assert published.status_code == 200
    assert published.json()["state"] == "published"

    client.post("/api/v1/auth/logout")
    login(client, "admin")
    other_educator = client.post(
        "/api/v1/admin/users",
        json={
            "email": "other-educator@example.edu",
            "full_name": "Other Educator",
            "password": "safe-password",
            "role": "educator",
        },
    )
    assert other_educator.status_code == 201
    client.post("/api/v1/auth/logout")
    response = client.post(
        "/api/v1/auth/login",
        json={
            "email": "other-educator@example.edu",
            "password": "safe-password",
        },
    )
    assert response.status_code == 200
    assert client.get(f"/api/v1/courses/{course['id']}").status_code == 403


def test_uploaded_and_linked_materials_remain_accessible_to_course_members(
    lms_context: tuple[TestClient, Session],
) -> None:
    client, session = lms_context
    login(client, "educator")
    course = client.get("/api/v1/courses").json()[0]
    module = client.get(f"/api/v1/courses/{course['id']}/modules").json()[0]

    document = Document()
    document.add_heading("Hadamard gates", level=1)
    document.add_paragraph("A Hadamard gate prepares an equal superposition before measurement.")
    content = BytesIO()
    document.save(content)
    response = client.post(
        f"/api/v1/courses/{course['id']}/materials/upload",
        params={"module_id": module["id"]},
        files={
            "file": (
                "hadamard-notes.docx",
                content.getvalue(),
                ("application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            )
        },
    )
    assert response.status_code == 201
    material = response.json()
    assert material["indexing_status"] == "indexed"

    downloaded = client.get(f"/api/v1/courses/{course['id']}/materials/{material['id']}/content")
    assert downloaded.status_code == 200
    assert downloaded.content == content.getvalue()
    assert "hadamard-notes.docx" in downloaded.headers["content-disposition"]
    assert downloaded.headers["x-content-type-options"] == "nosniff"

    retrieval = client.post(
        f"/api/v1/courses/{course['id']}/retrieval/search",
        json={"query": "How does a Hadamard gate prepare superposition?"},
    )
    assert retrieval.status_code == 200, retrieval.text
    assert retrieval.json()["found"] is True
    assert retrieval.json()["hits"][0]["material_id"] == material["id"]
    assert retrieval.json()["hits"][0]["source_label"].startswith("hadamard-notes.docx")

    linked = client.post(
        f"/api/v1/courses/{course['id']}/materials/link",
        json={
            "source_url": "https://example.edu/quantum-notes",
            "module_id": module["id"],
        },
    ).json()
    redirected = client.get(
        f"/api/v1/courses/{course['id']}/materials/{linked['id']}/content",
        follow_redirects=False,
    )
    assert redirected.status_code == 307
    assert redirected.headers["location"] == "https://example.edu/quantum-notes"

    client.post("/api/v1/auth/logout")
    login(client, "student")
    assert (
        client.get(f"/api/v1/courses/{course['id']}/materials/{material['id']}/content").status_code
        == 200
    )

    assert (
        session.scalar(
            select(func.count())
            .select_from(PlatformAuditEvent)
            .where(PlatformAuditEvent.action == "material.uploaded")
        )
        == 1
    )


def test_prerequisites_attempt_history_events_and_points_once(
    lms_context: tuple[TestClient, Session],
) -> None:
    client, session = lms_context
    login(client, "student")
    dashboard = client.get("/api/v1/students/me/dashboard").json()
    first, second = dashboard["tasks"][:2]
    assert first["access_status"] == "available"
    assert second["access_status"] == "locked"
    assert "prerequisite" in dashboard["recommendations"][0]["reason"].casefold()
    assert "unlock" in dashboard["recommendations"][0]["reason"].casefold()
    assert (
        session.scalar(
            select(func.count())
            .select_from(Recommendation)
            .where(Recommendation.is_active.is_(True))
        )
        >= 1
    )

    locked = client.put(
        f"/api/v1/students/me/tasks/{second['id']}/draft",
        json={"code": "Hadamard"},
    )
    assert locked.status_code == 423

    assert client.get(f"/api/v1/students/me/tasks/{first['id']}").status_code == 200
    empty_draft = client.get(
        f"/api/v1/students/me/tasks/{first['id']}/draft",
    )
    assert empty_draft.status_code == 200
    assert empty_draft.json() is None
    saved = client.put(
        f"/api/v1/students/me/tasks/{first['id']}/draft",
        json={"answer": "b"},
    )
    assert saved.status_code == 200
    restored = client.get(
        f"/api/v1/students/me/tasks/{first['id']}/draft",
    )
    assert restored.status_code == 200
    assert restored.json() == saved.json()
    submitted = client.post(
        f"/api/v1/students/me/tasks/{first['id']}/submissions",
        json={"answer": "b"},
    )
    assert submitted.status_code == 201
    attempt = submitted.json()
    assert attempt["status"] == "completed"
    assert attempt["points_awarded"] == first["points"]
    assert attempt["feedback_reference"] == attempt["id"]

    repeated = client.post(
        f"/api/v1/students/me/tasks/{first['id']}/submissions",
        json={"answer": "b"},
    ).json()
    assert repeated["attempt_number"] == 2
    assert repeated["points_awarded"] == 0
    assert session.scalar(select(func.count()).select_from(TaskPointAward)) == 1
    assert session.scalar(select(func.count()).select_from(SubmissionAttempt)) == 2
    assert {
        event_type for event_type in session.scalars(select(LearningEvent.event_type)).all()
    } >= {
        LearningEventType.TASK_VIEW,
        LearningEventType.DRAFT_SAVE,
        LearningEventType.SUBMISSION,
        LearningEventType.COMPLETION,
    }
    assert all(
        str(event.pseudonymous_user_id).startswith("v1_")
        for event in session.scalars(select(LearningEvent)).all()
    )
    assert (
        session.scalar(
            select(func.count())
            .select_from(PlatformAuditEvent)
            .where(PlatformAuditEvent.action == "progress.updated")
        )
        == 2
    )

    refreshed = client.get("/api/v1/students/me/dashboard").json()
    assert refreshed["summary"]["points"] == first["points"]
    assert refreshed["tasks"][1]["access_status"] == "available"

    client.post("/api/v1/auth/logout")
    login(client, "educator")
    assert (
        client.get(
            f"/api/v1/students/me/tasks/{first['id']}/draft",
        ).status_code
        == 403
    )


def test_all_six_task_types_accept_and_mark_correct_responses(
    lms_context: tuple[TestClient, Session],
) -> None:
    client, _ = lms_context
    login(client, "student")
    tasks = client.get("/api/v1/students/me/dashboard").json()["tasks"]
    payloads = {
        "multiple_choice": {"answer": "b"},
        "multiple_answer": {"answer": '["c", "a"]'},
        "short_answer": {"answer": "A Hadamard gate creates a superposition."},
        "code_explanation": {
            "answer": "Measurement turns the superposition into a classical result."
        },
        "code_completion": {"code": "circuit.h(0)"},
        "quantum_circuit": {
            "circuit": {
                "qubits": 1,
                "operations": [{"gate": "h", "targets": [0]}],
            }
        },
    }
    assert {task["task_type"] for task in tasks} == set(payloads)
    for task in tasks:
        if task["task_type"] == "quantum_circuit":
            invalid = client.post(
                f"/api/v1/students/me/tasks/{task['id']}/submissions",
                json={
                    "circuit": {
                        "qubits": 1,
                        "operations": [{"gate": "h", "targets": [2]}],
                    }
                },
            )
            assert invalid.status_code == 422
        response = client.post(
            f"/api/v1/students/me/tasks/{task['id']}/submissions",
            json=payloads[task["task_type"]],
        )
        assert response.status_code == 201
        assert response.json()["status"] == "completed"
        assert response.json()["score"] == 100


def test_reminders_monitoring_and_admin_lifecycle(
    lms_context: tuple[TestClient, Session],
) -> None:
    client, session = lms_context
    task = session.scalar(select(LearningTask).order_by(LearningTask.position.desc()))
    assert task is not None
    task.due_at = datetime.now(UTC) - timedelta(days=2)
    session.commit()

    login(client, "student")
    first = client.get("/api/v1/students/me/dashboard").json()
    second = client.get("/api/v1/students/me/dashboard").json()
    assert len(first["reminders"]) == 1
    assert len(second["reminders"]) == 1
    assert session.scalar(select(func.count()).select_from(Reminder)) == 1
    reminder_id = first["reminders"][0]["id"]
    assert (
        client.patch(f"/api/v1/students/me/reminders/{reminder_id}/read").json()["is_read"] is True
    )

    client.post("/api/v1/auth/logout")
    login(client, "educator")
    students = client.get("/api/v1/educator/students").json()
    assert students[0]["at_risk"] is False  # Not started is a distinct state.
    dashboard = client.get("/api/v1/educator/dashboard").json()
    assert len(dashboard["weekly_engagement"]) == 7
    assert "task_type_performance" in dashboard
    course_id = dashboard["courses"][0]["id"]

    client.post("/api/v1/auth/logout")
    login(client, "admin")
    created = client.post(
        "/api/v1/admin/users",
        json={
            "email": "new-student@example.edu",
            "full_name": "New Student",
            "password": "safe-password",
            "role": "student",
        },
    )
    assert created.status_code == 201
    user = created.json()
    assert user["student_profile_id"] is not None
    assert client.post(f"/api/v1/admin/users/{user['id']}/deactivate").json()["is_active"] is False
    assert client.post(f"/api/v1/admin/users/{user['id']}/reactivate").json()["is_active"] is True
    updated = client.put(
        "/api/v1/admin/settings",
        json={
            "at_risk_threshold": 75,
            "reminders_enabled": True,
            "llm_provider": "openai",
            "llm_model": "configured-model",
        },
    )
    assert updated.status_code == 200
    assert updated.json()["at_risk_threshold"] == 75
    archived = client.post(f"/api/v1/admin/courses/{course_id}/archive")
    assert archived.status_code == 200
    assert archived.json()["state"] == CourseState.ARCHIVED
    assert (
        session.scalar(
            select(func.count())
            .select_from(PlatformAuditEvent)
            .where(PlatformAuditEvent.action == "account.deactivated")
        )
        == 1
    )


def test_submission_attempts_are_immutable(
    lms_context: tuple[TestClient, Session],
) -> None:
    client, session = lms_context
    login(client, "student")
    task = client.get("/api/v1/students/me/dashboard").json()["tasks"][0]
    client.post(
        f"/api/v1/students/me/tasks/{task['id']}/submissions",
        json={"answer": "b"},
    )
    attempt = session.scalar(select(SubmissionAttempt))
    assert attempt is not None
    assert attempt.status is AttemptStatus.COMPLETED
    attempt.score = 0
    with pytest.raises(RuntimeError, match="immutable"):
        session.commit()
    session.rollback()


def test_demo_bootstrap_endpoint_is_explicit(
    tmp_path: Path,
) -> None:
    database_path = tmp_path / "empty.db"
    engine = create_db_engine(f"sqlite:///{database_path.as_posix()}")
    Base.metadata.create_all(engine)
    factory = create_session_factory(engine)
    with factory() as session:
        app = create_app()
        app.dependency_overrides[get_db] = lambda: session
        with TestClient(app, client=("198.51.100.10", 50000)) as remote_client:
            assert remote_client.post("/api/v1/admin/bootstrap-demo").status_code == 404
        with TestClient(app, client=("127.0.0.1", 50000)) as client:
            assert client.get("/api/v1/courses").status_code == 401
            response = client.post("/api/v1/admin/bootstrap-demo")
            assert response.status_code == 201
            assert len(response.json()["users"]) == 3
            assert (
                client.post(
                    "/api/v1/auth/login",
                    json={
                        "email": "student@quantumlearn.demo",
                        "password": DEMO_PASSWORD,
                    },
                ).status_code
                == 200
            )
        app.dependency_overrides.clear()
    Base.metadata.drop_all(engine)
    engine.dispose()
