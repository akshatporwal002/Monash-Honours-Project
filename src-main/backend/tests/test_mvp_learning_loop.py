from __future__ import annotations

from collections.abc import Generator
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import create_engine, event, func, select, update
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy.pool import StaticPool

from app.api.audit_dependencies import (
    get_feedback_audit_events,
    get_student_audit_tracker,
)
from app.api.feedback_dependencies import get_feedback_executor
from app.api.learning_event_dependencies import get_feedback_view_tracker
from app.api.routes.lms import get_lms_material_storage
from app.core.config import settings
from app.db.base import Base
from app.db.session import get_db_session
from app.main import create_app
from app.models import (
    Enrollment,
    EnrollmentStatus,
    FeedbackRecord,
    FeedbackStatus,
    JudgeDecision,
    JudgeEvaluation,
    JudgeEvaluationStatus,
    LearningMaterial,
    MaterialChunk,
    MaterialIndexStatus,
    Recommendation,
    SubmissionAttempt,
    TaskPointAward,
    User,
    WorkflowRun,
    WorkflowStage,
)
from app.schemas.audit import AuditEventCommand
from app.services.audit_events import FeedbackAuditEvents, NullStudentAuditTracker
from app.services.feedback.application import InProcessFeedbackExecutor
from app.services.feedback.runtime import build_feedback_pipeline_for_repository
from app.services.learning_events import NoOpFeedbackViewTracker
from app.services.lms import DEMO_PASSWORD, bootstrap_demo
from app.services.rag.storage import LocalFileStorage


@dataclass(frozen=True, slots=True)
class MvpTestContext:
    client: TestClient
    session_factory: sessionmaker[Session]
    educator_id: int
    student_id: int
    admin_id: int


class NoOpAuditSink:
    def record(self, _: AuditEventCommand) -> None:
        return None


@pytest.fixture
def mvp_context(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> Generator[MvpTestContext, None, None]:
    engine = create_engine(
        "sqlite+pysqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )

    @event.listens_for(engine, "connect")
    def enable_foreign_keys(dbapi_connection: Any, _: Any) -> None:
        cursor = dbapi_connection.cursor()
        cursor.execute("PRAGMA foreign_keys=ON")
        cursor.close()

    Base.metadata.create_all(engine)
    factory = sessionmaker(
        bind=engine,
        autoflush=False,
        expire_on_commit=False,
    )
    with factory() as session:
        bootstrap_demo(session)
        session.execute(update(Enrollment).values(status=EnrollmentStatus.WITHDRAWN))
        educator_id = _user_id(session, "educator@quantumlearn.demo")
        student_id = _user_id(session, "student@quantumlearn.demo")
        admin_id = _user_id(session, "admin@quantumlearn.demo")
        session.commit()

    monkeypatch.setattr(settings, "llm_api_key", None)
    monkeypatch.setattr(settings, "csrf_enabled", True)
    monkeypatch.setattr(settings, "research_enabled", False)

    app = create_app()
    material_storage = LocalFileStorage(
        tmp_path / "material-storage",
        settings.rag_max_file_bytes,
    )

    def override_database() -> Generator[Session, None, None]:
        with factory() as session:
            yield session

    audit_events = FeedbackAuditEvents(NoOpAuditSink())
    app.dependency_overrides[get_db_session] = override_database
    app.dependency_overrides[get_feedback_audit_events] = lambda: audit_events
    app.dependency_overrides[get_feedback_executor] = lambda: InProcessFeedbackExecutor(
        factory,
        build_feedback_pipeline_for_repository,
        audit_events=audit_events,
    )
    app.dependency_overrides[get_lms_material_storage] = lambda: material_storage
    app.dependency_overrides[get_feedback_view_tracker] = NoOpFeedbackViewTracker
    app.dependency_overrides[get_student_audit_tracker] = NullStudentAuditTracker

    try:
        with TestClient(app) as client:
            yield MvpTestContext(
                client=client,
                session_factory=factory,
                educator_id=educator_id,
                student_id=student_id,
                admin_id=admin_id,
            )
    finally:
        app.dependency_overrides.clear()
        Base.metadata.drop_all(engine)
        engine.dispose()


def test_canonical_mvp_learning_loop(mvp_context: MvpTestContext) -> None:
    client = mvp_context.client

    educator_headers = _login(client, "educator@quantumlearn.demo", DEMO_PASSWORD)
    assert _json(client.get("/api/v1/auth/me"), 200)["role"] == "educator"

    course = _json(
        client.post(
            "/api/v1/courses",
            headers=educator_headers,
            json={
                "code": "MVP-101",
                "title": "Quantum Learning MVP",
                "description": "A minimal source-grounded learning pathway.",
            },
        ),
        201,
    )
    module = _json(
        client.post(
            f"/api/v1/courses/{course['id']}/modules",
            headers=educator_headers,
            json={
                "title": "Superposition",
                "description": "Build and explain a one-qubit superposition.",
                "position": 1,
            },
        ),
        201,
    )
    outcome = _json(
        client.post(
            f"/api/v1/modules/{module['id']}/outcomes",
            headers=educator_headers,
            json={
                "title": "Explain a Hadamard gate",
                "statement": ("Explain how a Hadamard gate creates a measurable superposition."),
                "kind": "weekly",
                "week_number": 1,
                "position": 1,
            },
        ),
        201,
    )
    material_bytes = _course_material_docx()
    material = _json(
        client.post(
            f"/api/v1/courses/{course['id']}/materials/upload",
            headers=educator_headers,
            files={
                "file": (
                    "quantum-foundations.docx",
                    material_bytes,
                    ("application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
                )
            },
        ),
        201,
    )
    assert material["module_id"] is None
    assert material["indexing_status"] == "indexed"
    educator_content = client.get(
        (f"/api/v1/courses/{course['id']}/materials/{material['id']}/content")
    )
    assert educator_content.status_code == 200
    assert educator_content.content == material_bytes
    assert "quantum-foundations.docx" in educator_content.headers["content-disposition"]

    with mvp_context.session_factory() as session:
        stored_material = session.get(LearningMaterial, material["id"])
        assert stored_material is not None
        assert stored_material.course_id == course["id"]
        assert stored_material.module_id is None
        assert stored_material.indexing_status is MaterialIndexStatus.INDEXED
        chunks = list(
            session.scalars(
                select(MaterialChunk)
                .where(MaterialChunk.material_id == stored_material.id)
                .order_by(MaterialChunk.chunk_index)
            )
        )
        assert chunks
        chunk_id = chunks[0].id
        assert "Hadamard gate" in chunks[0].chunk_text

    generated_tasks = _json(
        client.post(
            f"/api/v1/courses/{course['id']}/generate-tasks",
            headers=educator_headers,
            json={
                "learning_outcome_id": outcome["id"],
                "task_count": 3,
                "task_types": [
                    "multiple_choice",
                    "short_answer",
                    "quantum_circuit",
                ],
            },
        ),
        201,
    )
    first_task, second_task, third_task = generated_tasks
    assert first_task["source_references"] == [chunk_id]
    assert second_task["prerequisite_task_ids"] == [first_task["id"]]
    assert third_task["prerequisite_task_ids"] == [second_task["id"]]
    assert "Hadamard gate" in first_task["prompt"]

    enrollment = _json(
        client.post(
            f"/api/v1/courses/{course['id']}/enrollments",
            headers=educator_headers,
            json={"student_id": mvp_context.student_id},
        ),
        201,
    )
    assert enrollment["student_id"] == mvp_context.student_id
    assert (
        _json(
            client.post(
                f"/api/v1/courses/{course['id']}/publish",
                headers=educator_headers,
            ),
            200,
        )["state"]
        == "published"
    )
    assert client.get("/api/v1/students/me/dashboard").status_code == 403

    student_headers = _login(client, "student@quantumlearn.demo", DEMO_PASSWORD)
    assert _json(client.get("/api/v1/auth/me"), 200)["role"] == "student"
    student_content = client.get(
        (f"/api/v1/courses/{course['id']}/materials/{material['id']}/content")
    )
    assert student_content.status_code == 200
    assert student_content.content == material_bytes
    assert client.get("/api/v1/educator/dashboard").status_code == 403
    assert client.get("/api/v1/admin/settings").status_code == 403
    assert (
        client.post(
            "/api/v1/courses",
            headers=student_headers,
            json={"code": "FORBIDDEN", "title": "Forbidden"},
        ).status_code
        == 403
    )

    dashboard = _json(client.get("/api/v1/students/me/dashboard"), 200)
    tasks = dashboard["tasks"]
    assert [task["id"] for task in tasks] == [
        first_task["id"],
        second_task["id"],
        third_task["id"],
    ]
    assert [task["access_status"] for task in tasks] == [
        "available",
        "locked",
        "locked",
    ]
    assert dashboard["recommendations"][0]["task_id"] == first_task["id"]
    assert "unlock" in dashboard["recommendations"][0]["reason"].casefold()
    with mvp_context.session_factory() as session:
        persisted = session.scalar(
            select(Recommendation).where(
                Recommendation.student_id == mvp_context.student_id,
                Recommendation.task_id == first_task["id"],
                Recommendation.is_active.is_(True),
            )
        )
        assert persisted is not None

    locked_draft = client.put(
        f"/api/v1/students/me/tasks/{second_task['id']}/draft",
        headers=student_headers,
        json={"answer": "Trying to skip ahead"},
    )
    assert locked_draft.status_code == 423

    assert client.get(f"/api/v1/students/me/tasks/{first_task['id']}").status_code == 200
    draft = _json(
        client.put(
            f"/api/v1/students/me/tasks/{first_task['id']}/draft",
            headers=student_headers,
            json={"answer": "b"},
        ),
        200,
    )
    assert draft["answer"] == "b"
    in_progress = _json(client.get("/api/v1/students/me/dashboard"), 200)
    assert in_progress["tasks"][0]["access_status"] == "in_progress"

    first_attempt = _json(
        client.post(
            f"/api/v1/students/me/tasks/{first_task['id']}/submissions",
            headers=student_headers,
            json={"answer": "b"},
        ),
        201,
    )
    assert first_attempt["attempt_number"] == 1
    assert first_attempt["status"] == "completed"
    assert first_attempt["score"] == 100
    assert first_attempt["points_awarded"] == 100
    assert first_attempt["feedback_reference"] == first_attempt["id"]

    without_csrf = client.post(f"/api/v1/submissions/{first_attempt['id']}/feedback")
    assert without_csrf.status_code == 403, without_csrf.text
    assert without_csrf.json()["error"]["code"] == "csrf_validation_failed"

    terminal_feedback = _json(
        client.post(
            f"/api/v1/submissions/{first_attempt['id']}/feedback",
            headers=student_headers,
        ),
        200,
    )
    assert terminal_feedback["status"] in {"validated", "fallback"}
    assert terminal_feedback["status"] == "validated"
    assert terminal_feedback["feedback"]["kind"] == "validated"
    assert len(terminal_feedback["feedback"]["sources"]) == 1
    assert terminal_feedback["feedback"]["sources"][0]["source_id"] == chunk_id
    assert terminal_feedback["feedback"]["sources"][0]["label"].startswith(
        "quantum-foundations.docx"
    )

    with mvp_context.session_factory() as session:
        workflow = session.scalar(
            select(WorkflowRun).where(WorkflowRun.submission_id == first_attempt["id"])
        )
        assert workflow is not None
        assert workflow.current_stage is WorkflowStage.COMPLETED
        released = session.scalar(
            select(FeedbackRecord).where(
                FeedbackRecord.submission_id == first_attempt["id"],
                FeedbackRecord.status == FeedbackStatus.ACCEPTED,
            )
        )
        assert released is not None
        assert released.id == terminal_feedback["feedback"]["feedback_id"]
        assert released.source_references == [chunk_id]
        judge = session.scalar(
            select(JudgeEvaluation).where(JudgeEvaluation.feedback_id == released.id)
        )
        assert judge is not None
        assert judge.evaluation_status is JudgeEvaluationStatus.VALID
        assert judge.reported_decision is JudgeDecision.PASS
        assert judge.decision is JudgeDecision.PASS

    second_attempt = _json(
        client.post(
            f"/api/v1/students/me/tasks/{first_task['id']}/submissions",
            headers=student_headers,
            json={"answer": "a"},
        ),
        201,
    )
    assert second_attempt["id"] != first_attempt["id"]
    assert second_attempt["attempt_number"] == 2
    assert second_attempt["status"] == "submitted"
    assert second_attempt["score"] == 40
    assert second_attempt["points_awarded"] == 0

    history = _json(
        client.get(f"/api/v1/students/me/tasks/{first_task['id']}/submissions"),
        200,
    )
    assert [
        (attempt["attempt_number"], attempt["answer"], attempt["score"]) for attempt in history
    ] == [(2, "a", 40), (1, "b", 100)]

    after_resubmission = _json(
        client.get("/api/v1/students/me/dashboard"),
        200,
    )
    assert after_resubmission["summary"] == {
        "completed_tasks": 1,
        "total_tasks": 3,
        "completion_percentage": 33,
        "average_score": 40,
        "points": 100,
        "level": 1,
        "next_level_points": 400,
    }
    assert [task["access_status"] for task in after_resubmission["tasks"]] == [
        "completed",
        "available",
        "locked",
    ]
    assert any(
        recommendation["task_id"] == second_task["id"]
        for recommendation in after_resubmission["recommendations"]
    )

    with mvp_context.session_factory() as session:
        stored_attempts = list(
            session.scalars(
                select(SubmissionAttempt)
                .where(
                    SubmissionAttempt.student_id == mvp_context.student_id,
                    SubmissionAttempt.task_id == first_task["id"],
                )
                .order_by(SubmissionAttempt.attempt_number)
            ).all()
        )
        assert [
            (attempt.attempt_number, attempt.answer, attempt.score) for attempt in stored_attempts
        ] == [(1, "b", 100), (2, "a", 40)]
        assert (
            session.scalar(
                select(func.count())
                .select_from(TaskPointAward)
                .where(TaskPointAward.task_id == first_task["id"])
            )
            == 1
        )

    educator_headers = _login(client, "educator@quantumlearn.demo", DEMO_PASSWORD)
    educator_dashboard = _json(client.get("/api/v1/educator/dashboard"), 200)
    educator_course = next(
        item for item in educator_dashboard["courses"] if item["id"] == course["id"]
    )
    assert educator_course["student_count"] == 1
    assert educator_course["progress_percentage"] == 33
    assert client.get("/api/v1/students/me/dashboard").status_code == 403
    assert client.get(f"/api/v1/submissions/{first_attempt['id']}/feedback").status_code == 200

    admin_headers = _login(client, "admin@quantumlearn.demo", DEMO_PASSWORD)
    assert _json(client.get("/api/v1/admin/settings"), 200)["passing_score"] == 70
    assert client.get("/api/v1/educator/dashboard").status_code == 403
    other_student = _json(
        client.post(
            "/api/v1/admin/users",
            headers=admin_headers,
            json={
                "email": "other-student@example.edu",
                "full_name": "Other Student",
                "password": "other-student-password",
                "role": "student",
            },
        ),
        201,
    )
    assert other_student["student_profile_id"] is not None

    _login(client, "other-student@example.edu", "other-student-password")
    assert client.get(f"/api/v1/courses/{course['id']}").status_code == 403
    assert (
        client.get(
            (f"/api/v1/courses/{course['id']}/materials/{material['id']}/content")
        ).status_code
        == 403
    )
    hidden_feedback = client.get(f"/api/v1/submissions/{first_attempt['id']}/feedback")
    assert hidden_feedback.status_code == 404
    assert hidden_feedback.json()["error"]["code"] == "feedback_not_found"


def _login(client: TestClient, email: str, password: str) -> dict[str, str]:
    client.cookies.clear()
    response = client.post(
        "/api/v1/auth/login",
        json={"email": email, "password": password},
    )
    assert response.status_code == 200, response.text
    session_cookie = client.cookies.get(settings.session_cookie_name)
    csrf_token = client.cookies.get(settings.csrf_cookie_name)
    assert session_cookie
    assert csrf_token
    return {
        settings.csrf_header_name: csrf_token,
        "Origin": settings.frontend_origin,
    }


def _json(response: Any, expected_status: int) -> Any:
    assert response.status_code == expected_status, response.text
    return response.json()


def _user_id(session: Session, email: str) -> int:
    user_id = session.scalar(select(User.id).where(User.email == email))
    assert user_id is not None
    return user_id


def _course_material_docx() -> bytes:
    document = Document()
    document.add_heading("Hadamard and measurement", level=1)
    document.add_paragraph(
        "A Hadamard gate maps the zero state to an equal superposition. "
        "Measurement returns zero or one with equal ideal probability."
    )
    output = BytesIO()
    document.save(output)
    return output.getvalue()
